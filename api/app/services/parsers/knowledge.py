import os
from uuid import UUID, uuid4
from typing import Any
import docx
# from pypdf import PdfReader # Moved to lazy import
from langchain_text_splitters import RecursiveCharacterTextSplitter
from google import genai
from google.genai import types
import httpx

from api.app.models.entities import KnowledgeChunk

class KnowledgeParser:
    def __init__(
        self,
        gemini_api_key: str | None = None,
        ollama_url: str | None = None,
        ollama_embed_model: str = "nomic-embed-text",
        gemini_embed_model: str = "gemini-embedding-001",
    ):
        self.gemini_api_key = gemini_api_key
        self.ollama_url = ollama_url
        self.ollama_embed_model = ollama_embed_model
        self.gemini_embed_model = gemini_embed_model
        self.splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=100,
            separators=["\n\n", "\n", " ", ""]
        )
        self.client = None
        if self.gemini_api_key:
            self.client = genai.Client(api_key=self.gemini_api_key)

    def _extract_docx_text(self, file_path: str) -> str:
        doc = docx.Document(file_path)

        blocks: list[str] = []

        paragraph_texts = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
        if paragraph_texts:
            blocks.extend(paragraph_texts)

        for table in doc.tables:
            row_texts: list[str] = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text and cell.text.strip()]
                if cells:
                    row_texts.append(" | ".join(cells))
            if row_texts:
                blocks.append("\n".join(row_texts))

        return "\n\n".join(blocks)

    def parse_and_chunk(self, file_path: str) -> list[dict[str, Any]]:
        extension = os.path.splitext(file_path)[1].lower()
        full_text = ""

        if extension == ".pdf":
            from pypdf import PdfReader
            reader = PdfReader(file_path)
            for page in reader.pages:
                full_text += (page.extract_text() or "") + "\n"
        elif extension == ".docx":
            full_text = self._extract_docx_text(file_path)
        elif extension in [".txt", ".md"]:
            with open(file_path, "r", encoding="utf-8") as f:
                full_text = f.read()
        
        if not full_text.strip():
            return []

        chunks = self.splitter.split_text(full_text)
        return [{"content": chunk, "index": i} for i, chunk in enumerate(chunks)]

    async def generate_embeddings(self, chunks: list[dict[str, Any]]) -> list[list[float]]:
        texts = [c["content"] for c in chunks]
        embeddings = []
        
        # 1. Use Gemini if API key is provided
        if self.client:
            batch_size = 10
            for i in range(0, len(texts), batch_size):
                batch = texts[i:i + batch_size]
                result = self.client.models.embed_content(
                    model=self.gemini_embed_model,
                    contents=batch,
                    config=types.EmbedContentConfig(output_dimensionality=768),
                )
                embeddings.extend([e.values for e in result.embeddings])
            return embeddings

        # 2. Use Ollama local model if URL is provided
        if self.ollama_url:
            async with httpx.AsyncClient() as client:
                for text in texts:
                    response = await client.post(
                        f"{self.ollama_url.rstrip('/')}/api/embeddings",
                        json={"model": self.ollama_embed_model, "prompt": text},
                        timeout=30.0
                    )
                    response.raise_for_status()
                    data = response.json()
                    embeddings.append(data["embedding"])
            return embeddings

        # 3. Fallback dummy
        return [[0.1] * 768 for _ in chunks]

    async def process_file(
        self, 
        file_path: str, 
        workspace_id: UUID, 
        knowledge_asset_id: UUID, 
        knowledge_version_id: UUID
    ) -> list[KnowledgeChunk]:
        chunk_data = self.parse_and_chunk(file_path)
        if not chunk_data:
            raise ValueError("Knowledge asset produced no extractable text chunks")
             
        embeddings = await self.generate_embeddings(chunk_data)
        
        knowledge_chunks = []
        for i, (data, embedding) in enumerate(zip(chunk_data, embeddings)):
            chunk = KnowledgeChunk(
                id=uuid4(),
                knowledge_version_id=knowledge_version_id,
                asset_version_id=knowledge_version_id,
                chunk_index=data["index"],
                content=data["content"],
                embedding=embedding,
            )
            knowledge_chunks.append(chunk)
            
        return knowledge_chunks
